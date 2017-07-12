from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import NoAlertPresentException
import unittest, time, re, os, xdrlib, sys, xlrd
import xlsxwriter as wx
from ReadConfigFile import Config
import sys
import ftplib
import globalfile
import docx
import subprocess
import urllib2

def main1():
    fileHandle = open('defect_number.txt', 'r')
    myList = fileHandle.read().splitlines()
    print myList
    releaseDate = time.strftime('%Y/%m/%d', time.localtime(time.time()))
    globalfile.workbook = wx.Workbook('Defect release Content_%s.xlsx' % (releaseDate))
    #time.sleep(20)
    worksheet1 = workbook.add_worksheet("Defect")
    
    for number in myList:
        categoryValue = driver.find_element_by_xpath("//*[@id='sups[%d]']/td[2]" % (supIndex)).text
        #print categoryValue
        worksheet1.write(supIndex + 3, 0, categoryValue, textFormat)
        formValue = driver.find_element_by_xpath("//*[@id='sups[%d]']/td[3]/a" % (supIndex)).text
        #print formValue
        worksheet1.write(supIndex + 3, 1, formValue, textFormat)
        fixIDValue = driver.find_element_by_xpath("//*[@id='supsUniqueId%d']" % (supIndex)).text
        #print fixIDValue
        worksheet1.write(supIndex + 3, 2, fixIDValue, textFormat)
        releaseValue = driver.find_element_by_xpath("//*[@id='sups[%d]']/td[9]" % (supIndex)).text
        #print releaseValue
        worksheet1.write(supIndex + 3, 3, releaseValue, textFormat)
        levelValue = driver.find_element_by_xpath("//*[@id='sups[%d]']/td[10]" % (supIndex)).text
        #print levelValue
        worksheet1.write(supIndex + 3, 4, levelValue, textFormat)

#main1()
        
def checkSupersede():
    fileHandle = open('machinetype.txt', 'r')
    myList = fileHandle.read().splitlines()
    print myList
    
    #time.sleep(20)
    
    oslist=['win2012r2', 'win2016', 'rhel6', 'rhel7', 'sles10', 'sles11', 'sles12', 'esxi5.0', 'esxi5.1', 'esxi5.5', 'esxi6.0' ]
    for ostypename in oslist:
        print ostypename
        for MTtype in myList:
            print MTtype
            subprocess.call('C:\\onecli0329\\OneCli.exe update acquire --mt %s --metaonly --ostype %s >> MT_%s_supersedecheckresult.txt'% (MTtype,ostypename,MTtype) , shell=True )
    
#checkSupersede()

def word33():

    from docx import Document
    from docx.shared import Inches
    document = Document()
    
    document.add_heading('Document Title', 0)
    
    p = document.add_paragraph('A plain paragraph having some ')
    
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')
    
    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )
    
    #document.add_picture('monty-truth.png', width=Inches(1.25))
    
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'

    document.add_page_break()
    document.save('demo333.docx')


def outlook():
    app= 'Outlook'
    olApp = win32com.client.gencache.EnsureDispatch("Outlook.Application")    
    mail=olook.CreateItem(win32com.client.constants.olMailItem)
    mail.Recipients.Add('dchi@lenovo.com')    
    subj = mail.Subject = 'defect'
    body = "test test"
    #body.insert(0,"%s\r\n" %subj)
    #body.append("\r\nTest1111111111111111")
    mail.Body = '\r\n'.join(body)
    mail.Send()
    print "send ok"
    
    

def my():
    
    # Get the ASCII number of a character
    number = ord(char)
    
    # Get the character given by an ASCII number
    char = chr(number)
    
    
def regexp():
    import re
    fileHandle = open('A.txt', 'r')
    mydata = fileHandle.read()
    myList = mydata.splitlines()
    #print len(myList)
    
    for item in myList:      
        m = re.findall('[a-z][A-Z]{3}[a-z][A-Z]{3}[a-z]', item)
        
        if len(m) != 0:
            #print "========"
            #print myList.index(item)
            #print myList[myList.index(item)-1]
            #print item
            #print len(m)
            print m
            '''
            for g in range(0 , len(m)):
                #print len(m)
                #print m[g]
                
                str1 = ''.join(m[g])
                #print "str1:"
                print str1
                #print "item.find(str1):"
                #print item.find(str1)
                finalresult = ""
                for p in range (-4,5):
                    
                    
                    myresult = myList[myList.index(item)-p][item.find(str1)+4:item.find(str1)+5]
                    finalresult = finalresult + myresult.strip()
                    print finalresult
                    checked = re.findall('[A-Z]{3}[a-z][A-Z]{3}', finalresult)
                    if len(checked) != 0:
                        #print checked
                        #print str1
                        for j in range (-4,5):
                            print myList[myList.index(item)-j]
                        #print myresult = myList[myList.index(item)-p]
            '''
            #print "YA"
    #print m


#regexp()

def findhtmlsource():

    driver = webdriver.Firefox()
    base_url = "http://rsl-ossweb20.labs.lenovo.com:9084/OssWeb/html/OneStopShopRedLogin.html"
    driver.get(base_url)
    '''
    print "----------Build UXSP package begin----------"
    driver.find_element_by_name("userId").clear()
    driver.find_element_by_name("userId").send_keys("dchi@lenovo.com")
    driver.find_element_by_id("contra").clear()
    driver.find_element_by_id("contra").send_keys("4656")
    driver.find_element_by_name("submit").click()
    time.sleep(20)
    driver.find_element_by_link_text("Build Tool").click()
    time.sleep(30)
    '''
    
    
    html_source = driver.page_source
    
    
    
    aa = html_source.encode('utf-8')
    
   # print aa
    
    fileHandle = open("My_test1.txt", 'w+')
    print "1"
    fileHandle.write(aa)
    print fileHandle.read()
    myList = fileHandle.read().splitlines()
    print "3"
    print myList
'''    
    for line3 in myList:
        if 'relation' in line:
        print line3
'''
#findhtmlsource()    
#findhtmlsource()



def Speedtest():

    driver = webdriver.Firefox()
    base_url = "http://beta.speedtest.net/"
    driver.get(base_url)
    time.sleep(5)
    #driver.find_elements_by_class_name("start-text").click()
    #driver.find_element_by_xpath("//html/body/table/tbody/tr[1]/td[1]/table/tbody/tr[2]/td[2]/table/tbody/tr[2]/td[2]/span/a").click()
    driver.find_element_by_xpath("//*[@id='container']/div[2]/div/div/div/div[3]/div[1]/div[1]/a/span").click()
    
    
    #html_source = driver.page_source
    
    releaseDate = time.strftime('%Y%m%d%H%M%S', time.localtime(time.time()))
    print releaseDate
    workbook = wx.Workbook('SPEED_TEST_%s.xlsx' % (releaseDate))
    worksheet = workbook.add_worksheet("RESULT")
    worksheet.set_column('A:A', 10)
    worksheet.set_column('B:B', 10)
    worksheet.set_column('C:C', 10)

    yellow = workbook.add_format(
        {'font_name': 'SimSun', 'align': 'left', 'valign': 'vcenter', 'bg_color': '#FFB5B5', 'font_size': 12,
         'bold': 1, 'top': 1, 'left': 1, 'right': 1, 'bottom': 1})
    worksheet.merge_range(0, 0, 0, 2, '',yellow)
    worksheet.merge_range(1, 0, 1, 2, '',yellow)
    worksheet.merge_range(2, 0, 2, 2, '',yellow)
    worksheet.write(2, 0, "Time:%s" % (releaseDate), yellow)
    worksheet.write(1, 0, "Location: My_Lab", yellow)
    worksheet.write(0, 0, "Project: DebbyChi", yellow)
    
    titleFormat = workbook.add_format(
        {'font_name': 'Verdana', 'font_color': '#4F5055', 'bg_color': '#E0E0E0', 'font_size': 9, 'bold': 1, 'top': 1,
         'left': 1, 'right': 1, 'bottom': 1})
    worksheet.write(3, 0, "PING", titleFormat)
    worksheet.write(3, 1, "Download", titleFormat)
    worksheet.write(3, 2, "Upload", titleFormat)
    
    #aa = html_source.encode('utf-8')
    
    #print aa
    print driver.current_url
    textFormat = workbook.add_format(
        {'font_name': 'Verdana', 'font_color': '#800000', 'font_size': 8, 'top': 1, 'left': 1, 'right': 1, 'bottom': 1})

    
    
    while (True):
        time.sleep(2)
        buildResult = driver.current_url
        time.sleep(2)
        if "result" in buildResult:
            
            pingValue = driver.find_element_by_xpath("//*[@id='ping-value']").text
            #print categoryValue
            worksheet.write(4, 0, pingValue, textFormat)
            downloadValue = driver.find_element_by_xpath("//*[@id='container']/div[2]/div/div/div/div[3]/div[1]/div[3]/div/div[3]/div/div/div/div/div[1]/div[2]/div[2]/span[1]").text
            #print categoryValue
            worksheet.write(4, 1, downloadValue, textFormat)
            uploadValue = driver.find_element_by_xpath("//*[@id='container']/div[2]/div/div/div/div[3]/div[1]/div[3]/div/div[3]/div/div/div/div/div[1]/div[3]/div[2]/span[1]").text
            #print categoryValue
            worksheet.write(4, 2, uploadValue, textFormat)
            print "Speed_Test_Completed"
            break

        
  
    
    
    
    
    
    
    workbook.close()








def pyurl():

    driver = webdriver.Firefox()
    base_url = "http://www.pythonchallenge.com/pc/def/linkedlist.php?nothing=12345"
    driver.get(base_url)
    time.sleep(5)
    
    
    while (True):
        time.sleep(2)
        buildResult = driver.current_url
        html_source = driver.page_source
        target_number = re.findall('\d{1,}', html_source)
        
        
        print "target_number_list"
        print target_number
        target_number = ''.join(target_number)
        print target_number
        if target_number == "":
            
            
            print "I found it"
            break
        else:
            next_url = ("http://www.pythonchallenge.com/pc/def/linkedlist.php?nothing=%s" % (target_number))
            driver.get(next_url)
  
    
pyurl()
















#Speedtest()